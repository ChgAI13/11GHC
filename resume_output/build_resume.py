from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("/Users/hugochen/Documents/创业/resume_output")
PHOTO_SRC = Path("/Users/hugochen/Desktop/微信图片_20260702181147_429_34.jpg")
PHOTO_CROP = OUT_DIR / "portrait_crop.jpg"
DOCX_OUT = OUT_DIR / "陈宏谷_京东JDYOUNG_采销方向_精美完善版简历.docx"

FONT_CN = "PingFang SC"
FONT_EN = "Aptos"
INK = "1B1F23"
MUTED = "5B6472"
ACCENT = "B0182D"
DEEP = "0B2545"
BLUE = "2E5E8C"
LIGHT_BLUE = "EEF4FA"
LIGHT_GRAY = "F6F8FA"
BORDER = "D6DEE8"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def twips(value_cm):
    return int(Cm(value_cm).twips)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name_cn=FONT_CN, name_en=FONT_EN):
    run.font.name = name_en
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name_en)
    r_fonts.set(qn("w:hAnsi"), name_en)
    r_fonts.set(qn("w:eastAsia"), name_cn)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=None, color=None, bold=None):
    style.font.name = FONT_EN
    style.font.size = Pt(size) if size else None
    style.font.color.rgb = rgb(color) if color else None
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_EN)
    r_fonts.set(qn("w:hAnsi"), FONT_EN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=BORDER, size="6", val="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_cm, indent_dxa=0, total_width_cm=None):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total_width = twips(total_width_cm if total_width_cm else sum(widths_cm))
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(indent_dxa))
    ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(twips(width)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(twips(widths_cm[idx])))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


def add_para(container, text="", size=9.2, color=INK, bold=False, italic=False, after=2, before=0, align=None, line=1.08):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_line(container, parts, after=2, before=0, align=None, line=1.08):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    for text, opts in parts:
        r = p.add_run(text)
        set_run_font(
            r,
            size=opts.get("size", 9.2),
            color=opts.get("color", INK),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p_pr = p._element.get_or_add_pPr()
    bdr = p_pr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        p_pr.append(bdr)
    bottom = bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BORDER)
    r = p.add_run(text)
    set_run_font(r, size=10.7, color=DEEP, bold=True)
    return p


def add_bullet(doc, text, level=0, after=1.2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.50 + 0.32 * level)
    p.paragraph_format.first_line_indent = Cm(-0.20)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(" " + text)
    set_run_font(r, size=8.7, color=INK)
    return p


def add_role_header(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(title)
    set_run_font(r1, size=9.4, color=INK, bold=True)
    r2 = p.add_run("  |  " + subtitle)
    set_run_font(r2, size=8.8, color=MUTED)
    return p


def crop_photo():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(PHOTO_SRC).convert("RGB")
    width, height = image.size
    crop_w = int(width * 0.70)
    crop_h = int(crop_w / 0.74)
    left = (width - crop_w) // 2
    top = int(height * 0.065)
    crop = image.crop((left, top, left + crop_w, top + crop_h))
    crop.save(PHOTO_CROP, quality=96)


def build_doc():
    crop_photo()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.12)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.38)
    section.right_margin = Cm(1.38)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    styles = doc.styles
    set_style_font(styles["Normal"], size=9.2, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(2)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    set_style_font(styles["List Bullet"], size=8.7, color=INK)

    header = doc.add_table(rows=1, cols=2)
    set_table_geometry(header, [14.1, 4.1], total_width_cm=18.2)
    for row in header.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            set_cell_borders(cell, val="nil")
            clear_cell(cell)

    left, right = header.rows[0].cells
    name = left.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    r = name.add_run("陈宏谷")
    set_run_font(r, size=22, color=DEEP, bold=True)
    r = name.add_run("  CHEN HONGGU")
    set_run_font(r, size=9.2, color=MUTED, bold=True)

    add_rich_line(
        left,
        [
            ("京东 JDYOUNG｜采销方向", {"size": 10.6, "color": ACCENT, "bold": True}),
            ("  经济学背景 / 数据分析 / 零售履约实践", {"size": 9.0, "color": MUTED}),
        ],
        after=3,
    )
    add_rich_line(
        left,
        [
            ("电话：", {"size": 8.8, "color": MUTED, "bold": True}),
            ("18225157197", {"size": 8.8, "color": INK}),
            ("    邮箱：", {"size": 8.8, "color": MUTED, "bold": True}),
            ("Chenhg2026@163.com", {"size": 8.8, "color": INK}),
            ("    所在地：", {"size": 8.8, "color": MUTED, "bold": True}),
            ("澳大利亚·布里斯班", {"size": 8.8, "color": INK}),
        ],
        after=2,
    )
    add_para(
        left,
        "定位：以经济学训练理解消费需求与市场竞争，以数据整理和一线履约经验支持品类运营、供应协同与商品经营决策。",
        size=8.8,
        color=INK,
        after=0,
        line=1.05,
    )

    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(PHOTO_CROP), width=Cm(3.05), height=Cm(4.12))

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    highlights = doc.add_table(rows=1, cols=3)
    set_table_geometry(highlights, [6.05, 6.05, 6.10], total_width_cm=18.2)
    h_data = [
        ("采销岗位理解", "关注选品、定价、库存与用户需求之间的联动，能够从消费者与平台效率两端看问题。"),
        ("数据分析基础", "熟悉 Excel 数据整理、透视分析与可视化，可将分散信息转化为可执行判断。"),
        ("项目推进能力", "独立规划课程平台，完成需求梳理、规则建库、模块设计与体验优化。"),
    ]
    for idx, cell in enumerate(highlights.rows[0].cells):
        set_cell_shading(cell, LIGHT_BLUE if idx != 1 else LIGHT_GRAY)
        set_cell_borders(cell, color=BORDER, size="5")
        set_cell_margins(cell, top=95, bottom=85, start=120, end=120)
        clear_cell(cell)
        add_para(cell, h_data[idx][0], size=8.9, color=DEEP, bold=True, after=1)
        add_para(cell, h_data[idx][1], size=7.9, color=INK, after=0, line=1.05)

    add_section_heading(doc, "教育背景")
    add_rich_line(
        doc,
        [
            ("昆士兰大学（The University of Queensland）", {"size": 9.4, "color": INK, "bold": True}),
            ("  |  经济学学士 Bachelor of Economics  |  预计 2028 年毕业", {"size": 8.8, "color": MUTED}),
        ],
        after=1.5,
    )
    add_para(
        doc,
        "核心课程：中级微观经济学、中级宏观经济学、计量经济学、成本收益分析、公共经济学、竞争政策、环境经济学、卫生经济学、市场营销、商业写作。",
        size=8.55,
        color=INK,
        after=1.5,
        line=1.05,
    )

    add_section_heading(doc, "核心项目经历")
    add_role_header(doc, "UQ Economics Planner｜产品负责人（个人项目）", "课程规划 / GPA 计算 / 毕业审核工具")
    add_para(
        doc,
        "项目背景：针对 UQ Economics 学生课程信息分散、GPA 计算繁琐、毕业要求查询复杂等痛点，规划并持续开发一站式课程规划平台。",
        size=8.65,
        color=INK,
        after=1.2,
        line=1.05,
    )
    project_points = [
        "负责需求分析、功能规划与迭代方向制定，围绕学生选课、成绩管理和毕业审核搭建产品框架。",
        "梳理经济学培养方案、课程先修关系及毕业规则，建立课程数据库和规则信息底座。",
        "规划 GPA Calculator、Course Planner、Graduation Checker 等核心模块，优化课程查询与规划流程。",
        "将分散的课程信息整合到统一平台，为后续课程推荐和智能规划功能预留数据基础。",
    ]
    for point in project_points:
        add_bullet(doc, point)

    add_role_header(doc, "商业分析项目", "市场数据整理 / 消费需求分析 / 优化建议")
    for point in [
        "收集行业及市场数据，使用 Excel 完成数据清洗、分类汇总、透视分析与可视化呈现。",
        "运用经济学理论分析消费者需求、市场竞争和商业案例，形成对业务问题的结构化判断。",
        "输出商业分析报告，并提出具备执行方向的运营或市场优化建议。",
    ]:
        add_bullet(doc, point)

    add_section_heading(doc, "实践经历")
    add_role_header(doc, "Uber Eats 即时零售运营实践（兼职）", "履约效率 / 商户沟通 / 客户体验")
    for point in [
        "高峰时段合理规划配送路线，平衡时效、路线成本与订单履约质量。",
        "与商户及客户保持高效沟通，及时处理订单信息差，保障履约流程顺畅。",
        "在快节奏、高不确定性的即时零售场景中提升执行力、时间管理与应变能力。",
    ]:
        add_bullet(doc, point)

    add_section_heading(doc, "专业技能")
    skills = doc.add_table(rows=3, cols=2)
    set_table_geometry(skills, [3.15, 15.05], total_width_cm=18.2)
    rows = [
        ("数据工具", "Excel：函数、数据透视表、图表制作、数据整理与基础可视化"),
        ("办公表达", "Word / PowerPoint、商业写作、报告结构化呈现"),
        ("业务分析", "商业分析、市场分析、消费者需求分析、成本收益分析"),
    ]
    for ridx, row in enumerate(skills.rows):
        label, detail = rows[ridx]
        for cidx, cell in enumerate(row.cells):
            set_cell_shading(cell, WHITE if cidx else LIGHT_GRAY)
            set_cell_borders(cell, color=BORDER, size="4")
            set_cell_margins(cell, top=65, bottom=65, start=100, end=100)
            clear_cell(cell)
        add_para(row.cells[0], label, size=8.2, color=DEEP, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(row.cells[1], detail, size=8.25, color=INK, after=0)

    add_section_heading(doc, "职业动机与自我评价")
    add_para(
        doc,
        "关注互联网零售与消费市场，对采销岗位保持长期兴趣。具备数据分析意识、商业敏感度和较强学习能力，能够快速理解业务需求并推动任务落地。希望在京东采销体系中持续积累商品经营、品类运营和跨团队协同经验，成长为兼具数据判断与业务执行力的零售人才。",
        size=8.7,
        color=INK,
        after=0,
        line=1.06,
    )

    doc.core_properties.title = "陈宏谷 京东 JDYOUNG 采销方向简历"
    doc.core_properties.subject = "Resume for JDYOUNG Procurement and Sales"
    doc.core_properties.author = "陈宏谷"
    doc.save(DOCX_OUT)


if __name__ == "__main__":
    build_doc()
    print(DOCX_OUT)
